from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Palette couleurs ──────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1A, 0x3A, 0x6E)
BLUE_MED   = RGBColor(0x2E, 0x6D, 0xA8)
BLUE_LIGHT = RGBColor(0xD6, 0xE4, 0xF0)
ORANGE     = RGBColor(0xE8, 0x7C, 0x1E)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_BG    = RGBColor(0xF5, 0xF5, 0xF5)
YELLOW_BG  = RGBColor(0xFF, 0xF3, 0xCD)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   kwargs.get('val',   'single'))
        border.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', 'FFFFFF'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = BLUE_DARK
    p.runs[0].font.size = Pt(15)
    p.runs[0].bold = True
    return p

def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = BLUE_MED
    p.runs[0].font.size = Pt(12)
    p.runs[0].bold = True
    return p

def heading3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = BLUE_MED
    p.runs[0].font.size = Pt(11)
    p.runs[0].bold = True
    return p

def body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
            run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def bullet(doc, text, bold_parts=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
            run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def add_table_header(table, headers, bg=BLUE_MED, fg=WHITE):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = fg
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
        set_cell_border(cell, color='FFFFFF')

def add_row(table, values, bg=None, center=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        if center:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bg:
            set_cell_bg(cell, bg)
    return row

def add_note(doc, text, bg=YELLOW_BG):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('📸 ' + text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x85, 0x60, 0x04)
    set_cell_bg(cell, bg)
    doc.add_paragraph()

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE COUVERTURE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
r = title.add_run('MSPR — Rapport technique')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = BLUE_DARK

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub1.add_run('Industrialisation & mise en production')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = BLUE_MED

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Substitution Avion → Train')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = BLUE_MED

caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = caption.add_run('Conteneurisation · Tests automatisés · CI/CD · Supervision temps réel')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

bloc = doc.add_paragraph()
bloc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = bloc.add_run('Bloc E6.3 — TPRE622 · RNCP36581 · DIA/DIADS 2025-2026')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

school = doc.add_paragraph()
school.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = school.add_run('EPSI Montpellier')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = BLUE_DARK

doc.add_paragraph()
doc.add_paragraph()

# Fiche projet
fiche = doc.add_table(rows=8, cols=4)
fiche.style = 'Table Grid'
fiche.alignment = WD_TABLE_ALIGNMENT.CENTER

fiche_data = [
    ('Client',     'ObRail Europe',            'Année',         '2025 – 2026'),
    ('École',      'EPSI Montpellier',          'Formation',     'B3 Dév. IA & Data Science'),
    ('Projet',     'MSPR — Mise en production IA', 'Localisation', 'Montpellier'),
    ('Bloc',       'E6.3 — TPRE622 (RNCP36581)', 'Périmètre',   'Industrialisation / Prod'),
    ('Livrables',  'App Docker · CI/CD · Tests · Monitoring · Rapport',
                                               'Stack',         'FastAPI · React · PostgreSQL · Docker · Prometheus · Grafana'),
    ('Sources',    'GTFS · SNCF · INSEE · OurAirports · Back on Track',
                                               'Tests',         'Pytest · Playwright · GitHub Actions'),
    ('Équipe',     'ALINGRIN Maxime · MARCIN Matthieu\nMONTAGNON Romain · TROUILLET StèveJohn',
                                               'Coach MSPR',    'PISSOT Laurent'),
    ('Version',    'V1.0',                     'Date',          '[ 03/07/2026 ]'),
]

for row_idx, (k1, v1, k2, v2) in enumerate(fiche_data):
    row = fiche.rows[row_idx]
    for col_idx, (val, is_key) in enumerate([(k1, True), (v1, False), (k2, True), (v2, False)]):
        cell = row.cells[col_idx]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if is_key:
            run.bold = True
            run.font.color.rgb = WHITE
            set_cell_bg(cell, BLUE_DARK)
        else:
            run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
            set_cell_bg(cell, BLUE_LIGHT)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SOMMAIRE
# ══════════════════════════════════════════════════════════════════════════════
toc_title = doc.add_paragraph()
r = toc_title.add_run('Sommaire')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = BLUE_DARK
toc_title.paragraph_format.space_after = Pt(12)

add_toc(doc)
doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. CONTEXTE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '1. Contexte du projet')
body(doc, "ObRail Europe est un observatoire ferroviaire européen (client fictif du fil rouge MSPR). Sa mission : identifier les vols court-courrier remplaçables par le train et chiffrer le CO₂ économisé par passager, dans l'esprit de la loi française de 2023 qui supprime les vols de moins de 600 km doublés d'une liaison ferroviaire. Lors de la MSPR précédente (TPRE512), l'équipe a livré un entrepôt de données ferroviaires harmonisé et une API REST à l'état de prototype. Ce socle fonctionne, mais il reste artisanal : déploiement manuel, aucun test automatisé, aucune supervision. C'est précisément ce que cette MSPR vient corriger.")

# ══════════════════════════════════════════════════════════════════════════════
# 2. OBJECTIF
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '2. Objectif de cette MSPR')
body(doc, "L'objectif de cette troisième MSPR est d'industrialiser et de mettre en production la solution. Concrètement, il s'agit de transformer le prototype en une application web complète, reproductible et supervisée. Sept objectifs opérationnels structurent le travail :")
bullet(doc, "**Industrialisation :** stabiliser le backend et conteneuriser toute l'application (API, frontend, base de données, monitoring) pour un déploiement reproductible.", bold_parts=True)
bullet(doc, "**Interface professionnelle :** un frontend ergonomique et accessible, utilisable par des partenaires non techniques (ONG, institutions, opérateurs).", bold_parts=True)
bullet(doc, "**Stratégie de tests :** tests unitaires et d'intégration sur le backend, tests end-to-end sur le frontend, avec une couverture suffisante pour la production.", bold_parts=True)
bullet(doc, "**Pipeline CI/CD :** automatiser les tests, la construction des images Docker et la mise à disposition de versions testables.", bold_parts=True)
bullet(doc, "**Supervision :** un monitoring temps réel de la disponibilité, de la latence, du taux d'erreurs et des logs, avec un tableau de bord opérationnel.", bold_parts=True)
bullet(doc, "**Conformité :** RGPD, accessibilité numérique (RGAA) et sécurité de base (validation des entrées, gestion des erreurs, logs).", bold_parts=True)
bullet(doc, "**Rapport technique :** le présent document, qui décrit l'ensemble de la démarche.", bold_parts=True)

# ══════════════════════════════════════════════════════════════════════════════
# 3. SPÉCIFICATIONS FONCTIONNELLES — USER STORIES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '3. Spécifications fonctionnelles — User Stories')
body(doc, "Les user stories formalisent les besoins des utilisateurs finaux. Elles constituent le point de départ de la conception du frontend et servent de base aux scénarios de tests end-to-end (E2E) Playwright. Chaque user story est tracée vers les critères d'acceptance et les fichiers de tests correspondants.")

heading2(doc, '3.1 Utilisateurs cibles')

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Profil', 'Description'])
add_row(t, ['Partenaire institutionnel', 'Commission européenne, ONG (Transport & Environnement) — besoin d\'analyse et de reporting'], bg=GREY_BG)
add_row(t, ['Opérateur ferroviaire', 'SNCF, DB, ÖBB — besoin de valoriser leurs corridors substituables'])
add_row(t, ['Équipe interne ObRail', 'Administrateurs supervisant la disponibilité du service'], bg=GREY_BG)
doc.add_paragraph()

heading2(doc, '3.2 User Stories')

us_data = [
    (
        'US-01 — Navigation globale',
        'En tant que partenaire institutionnel,\nJe veux naviguer librement entre les sections de l\'application,\nAfin d\'accéder rapidement à l\'information qui m\'intéresse sans me perdre.',
        ['La barre de navigation est visible sur toutes les pages.',
         'Les 4 liens (Accueil, Trajets, Prédiction, Monitoring) sont accessibles en un clic.',
         'La page active est visuellement identifiée.',
         'Un lien d\'accès rapide au contenu principal est disponible (accessibilité RGAA).'],
        'frontend/e2e/navigation.spec.js — 11 tests'
    ),
    (
        'US-02 — Consultation et filtrage des trajets ferroviaires',
        'En tant qu\'opérateur ferroviaire,\nJe veux consulter et filtrer la liste des corridors ferroviaires européens,\nAfin d\'identifier rapidement les trajets substituables desservis par mon réseau.',
        ['La liste des trajets s\'affiche avec les informations clés (origine, destination, distance, CO₂).',
         'Un filtre par gare d\'origine et de destination est disponible.',
         'Un bouton de réinitialisation restaure la liste complète.',
         'Un message explicite s\'affiche si aucun résultat ne correspond aux filtres.',
         'En cas d\'indisponibilité de l\'API, un message d\'erreur clair est affiché.'],
        'frontend/e2e/trajets.spec.js — 17 tests'
    ),
    (
        'US-03 — Prédiction de substituabilité avion → train',
        'En tant que partenaire institutionnel,\nJe veux soumettre un corridor (origine, destination, distance, type de train) pour obtenir une prédiction IA,\nAfin de savoir si ce vol peut être remplacé par le train et quantifier le gain CO₂ par passager.',
        ['Le formulaire valide les champs obligatoires avant soumission (distance > 0, type de train requis).',
         'Le résultat affiche clairement le verdict : substituable ou non substituable.',
         'Si substituable, le gain CO₂ en kg est affiché avec la comparaison visuelle avion vs train.',
         'Le formulaire peut être réinitialisé pour une nouvelle prédiction.',
         'Les champs vides bloquent la soumission avec un message d\'erreur explicite.'],
        'frontend/e2e/prediction.spec.js — 21 tests'
    ),
    (
        'US-04 — Surveillance de l\'état du service',
        'En tant que membre de l\'équipe interne ObRail,\nJe veux consulter l\'état de santé de l\'API et accéder aux outils de supervision,\nAfin de détecter rapidement tout incident et intervenir avant qu\'il n\'impacte les partenaires.',
        ['Le statut de l\'API (opérationnel / dégradé) est affiché en temps réel.',
         'Un bouton "Actualiser" permet de forcer un nouveau contrôle de santé.',
         'Les liens vers Grafana et Prometheus sont accessibles depuis l\'interface.',
         'Le tableau des endpoints disponibles est affiché avec leur état.',
         'En cas d\'API indisponible, le statut "dégradé" est clairement signalé.'],
        'frontend/e2e/monitoring.spec.js — 16 tests'
    ),
    (
        'US-05 — Vue d\'ensemble des indicateurs clés',
        'En tant que partenaire institutionnel,\nJe veux accéder dès la page d\'accueil aux indicateurs clés du projet ObRail,\nAfin de saisir immédiatement l\'impact environnemental des corridors analysés.',
        ['Les KPIs principaux sont affichés (nombre de corridors, % substituables, CO₂ moyen économisé).',
         'Les modules applicatifs sont présentés avec une description claire.',
         'La barre de santé API indique l\'état du service en temps réel.',
         'En cas d\'API indisponible, les statistiques affichent un état dégradé explicite.'],
        'frontend/e2e/home.spec.js — 17 tests'
    ),
]

for us_title, story, criterias, tests_ref in us_data:
    heading3(doc, us_title)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    for line in story.split('\n'):
        run = p.add_run(line + '\n')
        run.font.size = Pt(11)
        run.italic = True

    crit_title = doc.add_paragraph()
    r = crit_title.add_run("Critères d'acceptance :")
    r.bold = True
    r.font.size = Pt(11)
    crit_title.paragraph_format.space_after = Pt(2)

    for c in criterias:
        bullet(doc, c)

    test_p = doc.add_paragraph()
    test_p.paragraph_format.space_before = Pt(4)
    r1 = test_p.add_run("Tests associés : ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = test_p.add_run(tests_ref)
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLUE_MED
    doc.add_paragraph()

heading2(doc, '3.3 Matrice de traçabilité User Stories → Tests')

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['User Story', 'Suite E2E', 'Tests backend liés', 'Risque AMDEC'])
matrix = [
    ('US-01 Navigation',  'navigation.spec.js (11 tests)',  'test_health.py',       'R3'),
    ('US-02 Trajets',     'trajets.spec.js (17 tests)',     'test_trajets.py',      'R2, R4'),
    ('US-03 Prédiction',  'prediction.spec.js (21 tests)', 'test_predict_api.py',  'R1, R4, R5'),
    ('US-04 Monitoring',  'monitoring.spec.js (16 tests)', 'test_health.py',       'R3'),
    ('US-05 Accueil',     'home.spec.js (17 tests)',        'test_stats.py',        'R2, R3'),
]
for i, row_data in enumerate(matrix):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '4. Architecture globale de la solution')
body(doc, "La solution suit une architecture en services conteneurisés, chacun isolé dans son propre conteneur Docker et relié aux autres par un réseau interne géré par Docker Compose. Cette séparation permet de faire évoluer, tester ou redémarrer un service sans impacter les autres.")

heading2(doc, '4.1 Schéma d\'architecture')
add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Schéma d'architecture propre (draw.io / Excalidraw) montrant les 7 services, les ports et les flèches de flux.")

heading2(doc, '4.2 Rôle de chaque composant')

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['Service', 'Technologie', 'Port', 'Rôle'])
arch_rows = [
    ('Frontend',       'React 19 + Nginx',         '80',   'Interface web publique, consultation et visualisation'),
    ('API',            'FastAPI (Python 3.11)',      '8000', 'Accès aux données harmonisées + endpoints ML'),
    ('Base de données','PostgreSQL 16',             '5432', 'Persistance des corridors (schéma en étoile)'),
    ('Prometheus',     'Prometheus',                '9090', 'Collecte des métriques de l\'API toutes les 15 s'),
    ('Grafana',        'Grafana',                   '3000', 'Tableaux de bord de supervision (9 panneaux)'),
    ('pgAdmin',        'pgAdmin 4',                 '5050', 'Administration graphique de la base'),
    ('ETL / Dashboard','PySpark / Streamlit',       '—',    'Chargement & viz (profils Docker optionnels)'),
]
for i, row_data in enumerate(arch_rows):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

body(doc, "Le découplage frontend / API / base est volontaire : l'interface ne connaît que l'API, jamais la base directement. Cela limite la surface d'attaque et permet de remplacer un composant (ex. la base) sans toucher au reste.")

# ══════════════════════════════════════════════════════════════════════════════
# 5. BACKEND
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '5. Backend — API REST FastAPI')
body(doc, "Le backend est une API REST développée avec FastAPI (Python 3.11). FastAPI a été retenu plutôt que Flask car il génère automatiquement la documentation interactive (Swagger / OpenAPI), valide les entrées via Pydantic, et expose nativement des métriques exploitables par Prometheus — trois besoins directs du cahier des charges.")

heading2(doc, '5.1 Endpoints principaux')

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Endpoint', 'Méthode', 'Rôle'])
ep_rows = [
    ('/health',               'GET',  'État du service + vérification de la base (503 si dégradé)'),
    ('/trajets',              'GET',  'Liste filtrable des corridors (origine, destination, substituabilité, pagination)'),
    ('/trajets/{id}',         'GET',  'Détail d\'un corridor (404 si introuvable)'),
    ('/stats/volumes',        'GET',  'Agrégats : répartition jour/nuit, volumes par opérateur, CO₂'),
    ('/predict/substitution', 'POST', 'ML — le corridor est-il substituable avion→train ?'),
    ('/predict/co2_saved',    'POST', 'ML — CO₂ économisé par passager (kg)'),
    ('/metrics',              'GET',  'Métriques au format Prometheus'),
    ('/docs · /redoc',        'GET',  'Documentation interactive Swagger & ReDoc'),
]
for i, row_data in enumerate(ep_rows):
    bg = GREY_BG if i % 2 == 0 else None
    row = add_row(t, row_data, bg=bg)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

body(doc, "Les quatre endpoints exigés par le cahier des charges (/trajets, /trajets/{id}, /stats/volumes, /health) sont tous opérationnels. Les endpoints /predict relèvent du précédent MSPR ; ils sont présents dans le code et supervisés, mais ne sont pas le cœur de cette MSPR d'industrialisation.")

heading2(doc, '5.2 Gestion des erreurs & sécurité de base')
bullet(doc, "**Statuts HTTP explicites :** 200 (succès), 404 (trajet introuvable), 422 (entrée invalide, validée par Pydantic), 500 (erreur interne), 503 (base indisponible sur /health).", bold_parts=True)
bullet(doc, "**Requêtes SQL paramétrées** (liaison via SQLAlchemy text() + paramètres) : protection contre l'injection SQL.", bold_parts=True)
bullet(doc, "**CORS restreint** aux origines du frontend (localhost:80 et localhost:5173 en développement).", bold_parts=True)
bullet(doc, "**Logs applicatifs structurés** (horodatage, niveau, message) pour un diagnostic rapide ; aucune donnée nominative journalisée.", bold_parts=True)
doc.add_paragraph()
add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Page Swagger de l'API (http://localhost:8000/docs) montrant la liste des endpoints dépliée, avec un /trajets exécuté et sa réponse 200.")

heading2(doc, '5.3 Documentation automatique')
body(doc, "FastAPI génère la spécification OpenAPI sans code supplémentaire. Chaque endpoint est documenté (résumé, description, modèle de réponse Pydantic), ce qui donne une page Swagger directement utilisable par un évaluateur ou un partenaire technique pour tester l'API depuis le navigateur.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. FRONTEND
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '6. Frontend — Interface React accessible')
body(doc, "Le frontend est une application React 19 construite avec Vite, puis servie en production par Nginx. React a été choisi car il fait partie des outils recommandés par le client et permet une interface réactive sans rechargement de page. L'interface compte quatre pages, pensées pour un public non technique.")

heading2(doc, '6.1 Fonctionnalités par page')

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Page', 'Contenu'])
page_rows = [
    ('Accueil',    'Indicateurs clés (corridors, CO₂ moyen, % substituables), état du service, accès aux modules'),
    ('Trajets',    'Tableau filtrable des corridors (gare départ/arrivée, substituabilité), CO₂ économisé, badges visuels'),
    ('Prédiction', 'Formulaire interrogeant les endpoints ML /predict (autocomplétion des gares, résultat substituabilité + CO₂)'),
    ('Monitoring', 'État de santé de l\'API en direct, liens vers Grafana et Prometheus, liste des endpoints surveillés'),
]
for i, row_data in enumerate(page_rows):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Page d'accueil du frontend (http://localhost) avec les indicateurs clés et le bandeau d'état du service.")
add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Page Trajets avec un filtre appliqué (ex. origine = Paris) et le tableau de résultats.")

heading2(doc, '6.2 Accessibilité (RGAA)')
body(doc, "L'accessibilité n'est pas cosmétique : elle est implémentée dans le code. Les éléments suivants ont été vérifiés dans les composants React :")
bullet(doc, "Lien d'évitement « Aller au contenu principal » (skip-link) en début de page.")
bullet(doc, "Rôles ARIA (role=\"banner\", \"navigation\", \"status\", \"alert\") et points de repère sémantiques (header, nav, main, section).")
bullet(doc, "Libellés accessibles (aria-label, aria-labelledby) sur les liens, boutons et zones d'indicateurs ; icônes décoratives masquées (aria-hidden).")
bullet(doc, "Formulaires avec <label> associés à chaque champ (htmlFor), tableaux avec en-têtes scope=\"col\".")
bullet(doc, "Zones dynamiques annoncées (aria-live, role=\"alert\") pour le chargement et les erreurs.")
doc.add_paragraph()
add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Audit accessibilité (extension Lighthouse ou axe DevTools) sur la page d'accueil, montrant le score d'accessibilité.")

heading2(doc, '6.3 Tests end-to-end (Playwright)')
body(doc, "Le parcours utilisateur est validé par cinq fichiers de tests Playwright (home, navigation, trajets, prédiction, monitoring), soit environ 760 lignes. Playwright a été préféré à Cypress pour son intégration native dans la CI GitHub Actions et sa rapidité. Ces tests vérifient l'affichage des pages, la navigation, le filtrage des trajets et le formulaire de prédiction.")
body(doc, "Chaque fichier de test est directement dérivé d'une user story (voir section 3) : les critères d'acceptance deviennent les assertions Playwright.")

# ══════════════════════════════════════════════════════════════════════════════
# 7. DOCKER
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '7. Conteneurisation Docker')
body(doc, "Toute l'application est conteneurisée. Chaque service possède son propre Dockerfile, et un fichier docker-compose.yml orchestre l'ensemble. L'objectif du cahier des charges est respecté : un évaluateur relance toute la plateforme en une seule commande.")

heading2(doc, '7.1 Procédure de lancement')
body(doc, "Pour démarrer l'ensemble de la solution :")
p = doc.add_paragraph(style='No Spacing')
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("cp .env.example .env\ndocker compose up -d")
r.font.name = 'Courier New'
r.font.size = Pt(10)
body(doc, "L'application est accessible sur http://localhost dès que l'API répond sur /health.")

heading2(doc, '7.2 Images construites')
bullet(doc, "**API :** image python:3.11-slim, dépendances installées dans une couche séparée du code (build plus rapide), healthcheck intégré sur /health.", bold_parts=True)
bullet(doc, "**Frontend :** build multi-stage — Node 20 compile l'application React, puis seule l'image Nginx légère (alpine) embarque le résultat. Image finale réduite.", bold_parts=True)
bullet(doc, "**Base de données :** image officielle postgres:16-alpine, schéma initialisé automatiquement via database/init.sql.", bold_parts=True)
bullet(doc, "**Monitoring :** images officielles Prometheus et Grafana, configuration injectée par volumes.", bold_parts=True)

heading2(doc, '7.3 Orchestration & persistance')
body(doc, "Le fichier docker-compose.yml déclare sept services : db, api, frontend, prometheus, grafana, pgadmin, plus deux services optionnels (dashboard Streamlit, ETL PySpark) isolés dans des profils Docker pour ne pas alourdir le lancement standard. La persistance est garantie par trois volumes nommés (postgres_data, prometheus_data, grafana_data) : les données et tableaux de bord survivent à un redémarrage. Les dépendances entre services sont ordonnées (l'API attend que la base soit « saine » grâce au healthcheck pg_isready).")

# ══════════════════════════════════════════════════════════════════════════════
# 8. CI/CD
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '8. Pipeline CI/CD')
body(doc, "L'intégration continue est gérée par GitHub Actions (fichier .github/workflows/ci.yml). GitHub Actions a été retenu car le dépôt est hébergé sur GitHub : aucune infrastructure supplémentaire à maintenir, et les secrets sont gérés nativement. Le pipeline se déclenche à chaque push sur main ou sur une branche feat/**, ainsi que sur chaque pull request vers main.")

heading2(doc, '8.1 Étapes du pipeline')

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Job', 'Contenu', 'Condition'])
ci_rows = [
    ('test-backend', 'Installe les dépendances Python, lance pytest avec seuil de couverture ≥ 80 %, publie le rapport HTML', 'Toujours'),
    ('test-e2e',     'Installe Node 20 + Playwright (Chromium), lance les tests E2E, publie le rapport Playwright', 'Toujours'),
    ('build-docker', 'Construit les images Docker API et Frontend (tag = commit SHA + latest)', 'Si les 2 jobs de test réussissent'),
]
for i, row_data in enumerate(ci_rows):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

body(doc, "L'enchaînement est volontaire : aucune image Docker n'est construite tant que les tests backend ET E2E ne sont pas verts. Une régression bloque donc la livraison avant même la création des images.")

heading2(doc, '8.2 Gestion des secrets & variables')
bullet(doc, "Les variables sensibles (identifiants base, mots de passe) ne sont jamais dans le code : elles passent par un fichier .env, fourni en modèle non sensible .env.example.")
bullet(doc, "Le fichier .env réel est exclu du dépôt (.gitignore).")
bullet(doc, "En CI, les valeurs nécessaires sont injectées via les secrets GitHub Actions, pas écrites en clair dans le YAML.")
doc.add_paragraph()
add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Onglet « Actions » de GitHub montrant un pipeline vert (les 3 jobs réussis : test-backend, test-e2e, build-docker).")

# ══════════════════════════════════════════════════════════════════════════════
# 9. TESTS
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '9. Tests automatisés & couverture')
body(doc, "La stratégie de test suit la pyramide classique : beaucoup de tests unitaires rapides à la base, des tests d'intégration au milieu, quelques tests end-to-end au sommet. Une analyse AMDEC (modes de défaillance) documentée dans docs/stratégie de test.md a servi à prioriser l'effort sur les risques critiques : modèle ML (R1), intégrité base (R2), endpoints API (R3), validation des entrées (R4).")

heading2(doc, '9.1 Répartition des tests')

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['Niveau', 'Outil', 'Fichiers', 'Ce qui est vérifié'])
test_rows = [
    ('Unitaires',   'Pytest',              'test_co2_estimation, test_predict_logic',                                  'Logique d\'estimation CO₂, logique de prédiction (modèles mockés)'),
    ('Intégration', 'Pytest + TestClient', 'test_health, test_trajets, test_stats, test_legacy_endpoints, test_predict_api', 'Endpoints API, codes HTTP, gestion d\'erreurs (base mockée)'),
    ('End-to-end',  'Playwright',          'home, navigation, trajets, prediction, monitoring',                        'Parcours utilisateur réel dans le navigateur'),
]
for i, row_data in enumerate(test_rows):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

body(doc, "Les tests backend isolent l'API : la base PostgreSQL et les modèles ML sont remplacés par des mocks (fixtures dans conftest.py). Les tests ne dépendent donc d'aucun service externe et tournent à l'identique en local et en CI.")

heading2(doc, '9.2 Couverture de code')
body(doc, "La couverture est mesurée par pytest-cov sur les modules critiques (api/, scripts/predict.py). Le pipeline échoue automatiquement si elle passe sous 80 % (--cov-fail-under=80). La valeur documentée dans le dépôt est de 83,91 %, au-dessus du seuil exigé.")

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Indicateur', 'Valeur'])
cov_rows = [
    ('Seuil minimal imposé (barrière CI)',   '80 %'),
    ('Couverture documentée du dépôt',        '83,91 %'),
    ('Nombre de fichiers de tests backend',   '7 (≈ 1 070 lignes)'),
    ('Nombre de specs E2E',                   '5 (≈ 760 lignes)'),
    ('Total tests (backend + E2E)',            '191 tests'),
]
for i, row_data in enumerate(cov_rows):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("Note de transparence : ")
r.bold = True
r.font.size = Pt(11)
r2 = p.add_run("la valeur de 83,91 % est celle affichée par le dépôt (badge README et barrière CI). Elle n'a pas été recalculée au moment de la rédaction de ce rapport — la commande pytest reste exécutable en local pour la régénérer.")
r2.italic = True
r2.font.size = Pt(11)
doc.add_paragraph()

add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Rapport de couverture HTML (htmlcov/index.html) ouvert dans le navigateur, montrant le pourcentage global.")

# ══════════════════════════════════════════════════════════════════════════════
# 10. SUPERVISION
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '10. Supervision & monitoring')
body(doc, "La supervision repose sur le couple Prometheus + Grafana, standard du marché et recommandé par le client. L'API expose ses métriques sur /metrics (via prometheus-fastapi-instrumentator). Prometheus les collecte toutes les 15 secondes et les conserve 15 jours ; Grafana les affiche dans un tableau de bord provisionné automatiquement au démarrage.")

heading2(doc, '10.1 Métriques suivies')

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Métrique', 'Intérêt opérationnel'])
metrics_rows = [
    ('Disponibilité de l\'API',                 'Le service répond-il ? (up/down)'),
    ('Débit (requêtes/s par endpoint)',          'Charge réelle et endpoints les plus sollicités'),
    ('Latence P50 / P95 / P99 (ms)',            'Temps de réponse, détection des ralentissements'),
    ('Taux d\'erreurs (%) et erreurs 4xx/5xx',  'Détection d\'incidents et d\'anomalies'),
    ('Volume de prédictions ML',                 'Suivi de l\'usage des endpoints /predict'),
]
for i, row_data in enumerate(metrics_rows):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

body(doc, "Le tableau de bord ObRail compte neuf panneaux couvrant la disponibilité, le débit, la latence, les erreurs et l'activité ML. C'est l'outil de diagnostic de l'équipe interne en cas d'incident.")
add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Tableau de bord Grafana ObRail (http://localhost:3000) affichant les 9 panneaux avec des données — latence, taux d'erreurs et débit visibles.")
add_note(doc, "CAPTURE D'ÉCRAN À INSÉRER : Page « Targets » de Prometheus (http://localhost:9090/targets) montrant la cible obrail-api en état UP.")

heading2(doc, '10.2 Politique de logs')
body(doc, "L'API journalise chaque erreur serveur avec horodatage, niveau et message clair (ex. « Erreur GET /trajets : … »). Les logs sont consultables via docker compose logs api. Aucune donnée personnelle n'est journalisée, conformément au RGPD.")

# ══════════════════════════════════════════════════════════════════════════════
# 11. SÉCURITÉ RGPD ACCESSIBILITÉ
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '11. Sécurité, RGPD & accessibilité')

heading2(doc, '11.1 RGPD')
bullet(doc, "Aucune donnée personnelle traitée : les données sont des corridors ferroviaires issus de sources ouvertes (GTFS SNCF, INSEE, OurAirports, Back on Track).")
bullet(doc, "Aucun log nominatif : la politique de logs n'enregistre pas d'identité ni d'adresse.")
bullet(doc, "Données sources en open data : pas de collecte auprès d'utilisateurs.")

heading2(doc, '11.2 Sécurité technique')
bullet(doc, "Entrées validées par Pydantic (types, bornes) → rejet en 422 des requêtes mal formées.")
bullet(doc, "Requêtes SQL paramétrées → protection contre l'injection SQL.")
bullet(doc, "CORS restreint aux origines connues du frontend.")
bullet(doc, "Secrets hors du code, via .env exclu du dépôt et secrets GitHub Actions en CI.")

heading2(doc, '11.3 Limites de sécurité assumées')
body(doc, "Par honnêteté vis-à-vis du jury : l'API n'a pas d'authentification (données publiques, périmètre observatoire ouvert). Par ailleurs, les identifiants Grafana et pgAdmin sont définis dans la configuration de démonstration — ils devront être déplacés vers des secrets et renforcés avant une vraie mise en production exposée sur Internet. Ces points sont identifiés comme des actions de durcissement à mener.")

heading2(doc, '11.4 Accessibilité')
body(doc, "L'accessibilité RGAA est traitée au chapitre 6.2 : skip-link, rôles ARIA, libellés associés, tableaux structurés, zones dynamiques annoncées. L'interface vise un public non technique (institutions, ONG, opérateurs).")

# ══════════════════════════════════════════════════════════════════════════════
# 12. MAINTENANCE & ROLLBACK
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '12. Plan de maintenance & rollback')
body(doc, "La conteneurisation et la CI/CD rendent la maintenance simple et prévisible. Les principes retenus :")

heading2(doc, '12.1 Maintenance courante')
bullet(doc, "**Versionnement par image :** chaque build CI tague les images avec le SHA du commit, ce qui permet de retrouver et redéployer exactement une version donnée.", bold_parts=True)
bullet(doc, "**Mises à jour :** modifier le code → push → la CI rejoue tests et build automatiquement ; aucune étape manuelle.", bold_parts=True)
bullet(doc, "**Surveillance :** Grafana et les alertes visuelles permettent de repérer une dégradation (latence, erreurs) avant qu'elle ne devienne critique.", bold_parts=True)
bullet(doc, "**Sauvegarde :** le volume postgres_data contient les données ; une sauvegarde régulière (pg_dump) est recommandée avant toute opération sensible.", bold_parts=True)

heading2(doc, '12.2 Procédure de rollback')

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Situation', 'Action de retour arrière'])
rollback_rows = [
    ('Une nouvelle version casse l\'API',    'Redéployer l\'image taguée du commit précédent (docker compose up avec le tag SHA antérieur)'),
    ('Un commit défectueux est passé',        'git revert du commit → la CI reconstruit une image saine automatiquement'),
    ('Migration de base problématique',       'Restaurer le dernier pg_dump dans le volume postgres_data'),
    ('Tableau de bord corrompu',              'Re-provisionnement Grafana au redémarrage (configuration versionnée dans le dépôt)'),
]
for i, row_data in enumerate(rollback_rows):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

body(doc, "Comme les tests bloquent la construction des images en cas de régression, la probabilité de devoir faire un rollback est déjà fortement réduite en amont.")

# ══════════════════════════════════════════════════════════════════════════════
# 13. LIMITES & PERSPECTIVES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '13. Limites actuelles & perspectives')
body(doc, "Ce rapport décrit l'état réel du projet. Plusieurs points restent à consolider :")
bullet(doc, "**Modèles ML :** les artefacts .joblib ne sont pas versionnés dans le dépôt (ils sont générés par les scripts d'entraînement et ignorés par git). Le conteneur API attend un dossier models/ ; en l'absence d'entraînement préalable, les endpoints /predict ne sont pas fonctionnels. C'est cohérent car le volet IA relève de la MSPR TPRE622.", bold_parts=True)
bullet(doc, "**Service ETL :** le profil etl du docker-compose référence un chemin de données spécifique à une machine (montage local) ; il faudra le rendre générique pour un autre évaluateur.", bold_parts=True)
bullet(doc, "**Sécurité :** absence d'authentification et secrets de démonstration à durcir (voir 11.3) avant une exposition publique réelle.", bold_parts=True)
bullet(doc, "**Déploiement :** la CI construit les images mais ne déploie pas encore vers un environnement distant (pas de registre / serveur cible) ; la « mise à disposition » reste locale.", bold_parts=True)

# ══════════════════════════════════════════════════════════════════════════════
# 14. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '14. Conclusion')
body(doc, "Cette MSPR a transformé un prototype artisanal en une application web complète, conteneurisée, testée et supervisée. Les sept objectifs du cahier des charges sont couverts : industrialisation Docker, frontend accessible, stratégie de tests avec couverture mesurée, pipeline CI/CD automatisé, supervision Prometheus/Grafana, conformité RGPD/RGAA, et le présent rapport. L'évaluateur peut relancer toute la plateforme en une seule commande.")

# ══════════════════════════════════════════════════════════════════════════════
# 15. ANNEXES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, '15. Annexes')

heading2(doc, '15.1 Glossaire')

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Terme', 'Définition'])
gloss = [
    ('API REST',           'Interface permettant à des applications de communiquer via HTTP (requêtes / réponses JSON)'),
    ('Conteneur / Docker', 'Environnement isolé embarquant une application et ses dépendances, reproductible partout'),
    ('CI/CD',              'Intégration et livraison continues : tests et builds automatisés à chaque changement de code'),
    ('Pyramide de tests',  'Stratégie : beaucoup d\'unitaires, moins d\'intégration, peu d\'E2E'),
    ('E2E (end-to-end)',   'Test du parcours utilisateur complet dans un vrai navigateur'),
    ('AMDEC',              'Analyse des modes de défaillance pour prioriser l\'effort de test selon la criticité'),
    ('Prometheus / Grafana','Outils de collecte de métriques et d\'affichage de tableaux de bord'),
    ('RGAA',               'Référentiel français d\'accessibilité des sites web'),
    ('Substituabilité',    'Capacité d\'un corridor ferroviaire à remplacer un vol court-courrier'),
    ('User Story',         'Expression d\'un besoin utilisateur sous la forme "En tant que… Je veux… Afin de…"'),
]
for i, row_data in enumerate(gloss):
    bg = GREY_BG if i % 2 == 0 else None
    add_row(t, row_data, bg=bg)
doc.add_paragraph()

heading2(doc, '15.2 Références & dépôt')
bullet(doc, "Dépôt du projet : code source complet (backend, frontend, Docker, CI, tests, monitoring).")
bullet(doc, "Documentation technique : README.md (démarrage, URLs, structure), docs/stratégie de test.md (AMDEC, user stories, plans de test).")
bullet(doc, "Sources de données : GTFS SNCF, INSEE, OurAirports, Back on Track (open data).")
bullet(doc, "Méthode CO₂ : EcoPassenger (UIC / IFEU).")

# ── Sauvegarde ────────────────────────────────────────────────────────────────
out = '/Users/montagnonromain/Dossier perso Locaux/epsi-cours/MSPR/docs/Rapport_Technique_MSPR_E6.3_ObRail.docx'
doc.save(out)
print(f'Généré : {out}')
